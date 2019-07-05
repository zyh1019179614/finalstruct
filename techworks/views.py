from django.shortcuts import render, HttpResponse
from django.http.response import JsonResponse
from django.views import View
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt

from .models import WorkInfo, Appendix
from django.db.models import Q
from competition.models import CompetitionRegistration, Competition
from users.models import Student
from django.db import transaction
from competition.views import GetUserIdentitiy
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
import datetime, time, random
import json
import docx
from docx import Document
import re
import convertapi
from docx.oxml.ns import qn

from docx.shared import Pt


# Create your views here.
@csrf_exempt
def work_info(request):
    DEGREE = ['大专','大学本科','硕士研究生','博士研究生']
    user_name, user_identity = GetUserIdentitiy(request)
    information = {}
    workInfo = WorkInfo.objects.get(id=request.GET['work_id'])
    workAppendix = Appendix.objects.filter(work_id=request.GET['work_id'])
    workAuthor = CompetitionRegistration.objects.get(id=request.GET['work_id'])
    workFirstAuthorInfo = Student.objects.get(user_id=workAuthor.first_auth_id)
    authorList = []

    if(workAuthor.second_auth_id):
        workSecondAuthorInfo = Student.objects.get(user_id=workAuthor.second_auth_id)
        authorList.append({'name': workSecondAuthorInfo.name,
                           'stu_id': workSecondAuthorInfo.stu_id,
                           'degree': DEGREE[workSecondAuthorInfo.degree],
                           'phone': workSecondAuthorInfo.phone,
                           'email': workSecondAuthorInfo.user.email})
    else:
        workSecondAuthorInfo = ''

    if (workAuthor.third_auth_id):
        workThirdAuthorInfo = Student.objects.get(user_id=workAuthor.third_auth_id)
        authorList.append({'name': workThirdAuthorInfo.name,
                           'stu_id': workThirdAuthorInfo.stu_id,
                           'degree': DEGREE[workThirdAuthorInfo.degree],
                           'phone': workThirdAuthorInfo.phone,
                           'email': workThirdAuthorInfo.user.email})
    else:
        workThirdAuthorInfo = ''

    if (workAuthor.forth_auth_id):
        workForthAuthorInfo = Student.objects.get(user_id=workAuthor.forth_auth_id)
        authorList.append({'name': workForthAuthorInfo.name,
                           'stu_id': workForthAuthorInfo.stu_id,
                           'degree': DEGREE[workForthAuthorInfo.degree],
                           'phone': workForthAuthorInfo.phone,
                           'email': workForthAuthorInfo.user.email})
    else:
        workFourthAuthorInfo = ''

    if (workAuthor.fifth_auth_id):
        workFifthAuthorInfo = Student.objects.get(user_id=workAuthor.fifth_auth_id)
        authorList.append({'name': workFifthAuthorInfo.name,
                           'stu_id': workFifthAuthorInfo.stu_id,
                           'degree': DEGREE[workFifthAuthorInfo.degree],
                           'phone': workFifthAuthorInfo.phone,
                           'email': workFifthAuthorInfo.user.email})
    else:
        workFifthAuthorInfo = ''

    information = {'workInfo': workInfo}
    information['workAppendix']= workAppendix
    information['workAuthor'] = workAuthor
    information['authorList'] = authorList
    information['username'] = user_name
    information['useridentity'] = user_identity

    return render(request, '../templates/viewWorkInfo.html', information)

def searchstu(request):
    DEGREE_MAP = {
        1: "大专",
        2: "大学本科",
        3: "硕士研究生",
        4: "博士研究生",
    }
    stuid = request.GET.get('stu_id')
    student = Student.objects.get(stu_id=stuid)
    ret = {"student": {
        "stu_id": student.stu_id,
        "name": student.name,
        "degree": DEGREE_MAP[student.degree],
        "phone": student.phone,
        "email": student.user.email
    }}
    return JsonResponse(ret)

def deletework(request):
    work_id = request.GET.get('work_id')
    work = WorkInfo.objects.filter(work_id=work_id)
    work.delete()
    return HttpResponse(status=200)

# @login_required(login_url='/login/')
# @method_decorator(login_required,name='dispatch')
class TechWorkListView(View):
    # @login_required(login_url='/login/')
    def get(self, request):
        # if request.GET.get('workid', None) is None:
        #     WorkInfo.objects.create()
        username = request.user.username
        if username == "":
            username = "16211086"
        registrationlist = CompetitionRegistration.objects.filter(Q(first_auth__user__username=username)
                                                                  | Q(second_auth__user__username=username)
                                                                  | Q(third_auth__user__username=username)
                                                                  | Q(forth_auth__user__username=username)
                                                                  | Q(fifth_auth__user__username=username)
                                                                  )
        worklist_origin = WorkInfo.objects.filter(registration__in=registrationlist)
        # work = WorkInfo.objects.get(work_id='1').wo
        WORKTYPE_MAP = {
            1: "科技发明制作",
            2: "调查报告和学术论文"
        }
        FIELD_MAP = {
            1: "A",
            2: "B",
            3: "C",
            4: "D",
            5: "E",
            6: "F",
        }

        worklist_ret = []
        for work in worklist_origin:
            worklist_ret.append({
                'work_id': work.work_id,
                'first_author': work.registration.first_auth.name,
                'title': work.title,
                'work_type': WORKTYPE_MAP.get(work.work_type, 1),
                'field': FIELD_MAP.get(work.field, 1),
                'competition': work.registration.competition.title,
                'submitted': "已提交" if work.submitted else "暂存",
            })

        paginator = Paginator(worklist_ret, 10)  # 每页10条
        total = len(worklist_ret)
        list = paginator.page(1)
        page = request.GET.get('page')
        try:
            list = paginator.page(page)  # contacts为Page对象！
        except PageNotAnInteger:
            # If page is not an integer, deliver first page.
            list = paginator.page(1)
        except EmptyPage:
            # If page is out of range (e.g. 9999), deliver last page of results.
            list = paginator.page(paginator.num_pages)

        user_name, user_identity = GetUserIdentitiy(request)
        return render(request, 'techwork_list.html', {'worklist': list.object_list,'useridentity':user_identity,
                                                      'username':user_name, 'num':list.number,'total':total})


class TechWorkView(View):
    def get(self, request):
        work_id = request.GET.get('workid', None)
        comptition_id = request.GET.get('cptid', None)
        username = request.user.username
        if username == "":
            username = "16211086"
        work = None
        company_ret = []
        show_list = []
        invest_list = []
        file_docu = []
        file_photo = []
        file_video = []
        if work_id is None:
            if comptition_id is not None:
                registration = CompetitionRegistration.objects.create(
                    first_auth=Student.objects.get(user__username=username),
                    competition=Competition.objects.get(id=comptition_id))
                registration.save()
                work_cnt = int(time.time())*100 + random.randint(0, 10000)
                work = WorkInfo.objects.create(registration=registration, title="", detail="", innovation="",
                                               keywords="", avg_score=0, work_id=work_cnt, work_type=1,
                                               field=1)
                work.save()
            else:
                return HttpResponse(status=400)
        else:
            work = WorkInfo.objects.get(work_id=work_id)
            registration = work.registration
            company = []
            if registration.second_auth is not None:
                auth = registration.second_auth

                company.append(registration.second_auth)
            if registration.third_auth is not None:
                company.append(registration.third_auth)
            if registration.forth_auth is not None:
                company.append(registration.forth_auth)
            if registration.fifth_auth is not None:
                company.append(registration.fifth_auth)
            company_ret = []
            for auth in company:
                company_ret.append({"stu_id": auth.stu_id,
                                    "name": auth.name,
                                    "degree": auth.degree,
                                    "phone": auth.phone,
                                    "email": auth.user.email})
            if work.work_type == 1:
                show_list = json.loads(work.labels)['labels']
            else:
                invest_list = json.loads(work.labels)['labels']

            filelist = Appendix.objects.filter(work__work_id=work.work_id)

            for file in filelist:
                if (file.appendix_type == 0):
                    file_docu.append({"name": file.filename, "url": file.file.url})
                if (file.appendix_type == 1):
                    file_photo.append({"name": file.filename, "url": file.file.url})
                if (file.appendix_type == 2):
                    file_video.append({"name": file.filename, "url": file.file.url})

        user_name, user_identity = GetUserIdentitiy(request)
        return render(request, 'submit_techwork.html', {'work': work, 'company': company_ret,
                                                        'show_list': show_list, 'invest_list':invest_list,
                                                        'docu': file_docu, 'photo': file_photo, 'video': file_video,
                                                        'username':user_name, 'useridentity':user_identity})

    def post(self, request):
        try:
            with transaction.atomic():
                work_id = int(request.POST.get('work_id'))
                work = WorkInfo.objects.get(work_id=work_id)
                registration = work.registration
                AUTH_MAP = {
                    0: registration.second_auth,
                    1: registration.third_auth,
                    2: registration.forth_auth,
                    3: registration.fifth_auth
                }
                team = json.loads(request.POST.get("company_table", []))
                for id in range(0, 4):
                    if len(team) > id:
                        stu = team[id]
                    else:
                        stu = None
                    if id == 0:
                        registration.second_auth = Student.objects.get(stu_id=stu["stu_id"]) if stu else None
                    if id == 1:
                        registration.third_auth = Student.objects.get(stu_id=stu["stu_id"]) if stu else None
                    if id == 2:
                        registration.forth_auth = Student.objects.get(stu_id=stu["stu_id"]) if stu else None
                    if id == 3:
                        registration.fifth_auth = Student.objects.get(stu_id=stu["stu_id"]) if stu else None

                work.title = request.POST.get('title', work.title)
                work.detail = request.POST.get('detail', work.detail)
                work.innovation = request.POST.get('innovation', work.innovation)
                work.keywords = request.POST.get('keywords', work.keywords)
                work.work_type = int(request.POST.get('work_type', work.work_type))
                work.field = int(request.POST.get("field", work.field))

                if work.work_type == 1:
                    label_list = json.loads(request.POST.get('show_type', work.labels))
                else:
                    label_list = json.loads(request.POST.get('invest_type', work.labels))

                work.labels = json.dumps({"labels":label_list})

                work.submitted = True if int(request.POST.get("submitted", work.submitted)) == 1 else False
                FILE_TYLE_MAP = {
                    "document": 0,
                    "photo": 1,
                    "video": 2
                }

                deleteList = json.loads(request.POST.get('deleteList', "{\"deletelist\":[]}"))["deletelist"]
                appendixList = Appendix.objects.filter(work=work)
                for deletefile in deleteList:
                    for appdendix in appendixList:
                        if appdendix.file.url == deletefile:
                            appdendix.delete()

                if len(request.FILES) > 0:
                    for k, file in request.FILES.items():
                        # file = file[0]
                        filename = file._name
                        filedata = file
                        filetype = FILE_TYLE_MAP.get(k.split('_')[0], 0)
                        filedate = datetime.date.today()
                        Appendix.objects.create(filename=filename, appendix_type=filetype, upload_date=filedate,
                                                file=filedata, work=work)

                # team = json.loads(request.POST.get('company_table')) if request.POST.get(
                #     'company_table') != "" else None
                #
                # if team:
                #     for order, auth in enumerate(team):
                #         AUTH_MAP[order] = Student.objects.get(stu_id=auth['stu_id'])

                registration.save()
                work.save()
        except Exception as e:
            print(e)
            return HttpResponse(status=400)
        return HttpResponse(status=200)

        # file = request.FILES.get("file")
        # dataPOST = json.loads(request.body)
        # dataPOST['birthdate'] = dataPOST['birthdate'][0:10]
        # dataPOST['enroll_time'] = dataPOST['enroll_time'][0:10]
        # register_form = RegisterForm(dataPOST)
        # if register_form.is_valid():
        #     email = dataPOST.get('email', '')
        #     if Student.objects.filter(user__email=email):
        #         return render(request, "register.html", {"register_form": register_form, "msg": "该邮箱已注册"})
        #     password = dataPOST.get('password', '')
        #
        #     # 创建BaseUser对象
        #     baseuser = BaseUser()
        #     baseuser.username = dataPOST.get('stu_id', '')
        #     baseuser.password = make_password(password)
        #     baseuser.type = 1
        #     baseuser.email = email
        #     baseuser.save()



def generatePdf(request):
    if request.method == 'POST':
        workid = request.POST.get('workid')
        path = 'static/pdf/emptytable.docx'
        document = Document(path)  # 读入文件
        work_info = WorkInfo.objects.get(work_id=workid)
        registration = work_info.registration
        # 获取段落文字
        para = document.paragraphs
        para[0].text = re.sub("（系统自动生成）", workid, para[0].text)
        para[0].font
        para[8].text = re.sub("（名称）", work_info.title, para[8].text)
        if (work_info.work_type == 1):
            para[12].text = re.sub("□", "☑", para[12].text)
        elif (work_info.work_type == 2):
            para[13].text = re.sub("□", "☑", para[13].text)

        tables = document.tables  # 获取文件中的表格集
        table = tables[0]  # 获取文件中的第一个表格
        leng = len(table.rows)
        company = []

        if registration.first_auth is not None:
            auth = registration.first_auth
        if registration.second_auth is not None:
            company.append(registration.second_auth)
        if registration.third_auth is not None:
            company.append(registration.third_auth)
        if registration.forth_auth is not None:
            company.append(registration.forth_auth)
        if registration.fifth_auth is not None:
            company.append(registration.fifth_auth)

        # 表格1
        table.cell(0, 3).text = auth.name
        table.cell(0, 6).text = auth.stu_id
        table.cell(0, 8).text = auth.birthdate.strftime("%Y-%m-%d")
        # 替换学历

        degreechange = r"（  ）"
        if (auth.degree == 1):
            table.cell(1, 3).text = re.sub(degreechange, "（A）", table.cell(1, 3).text)
        elif (auth.degree == 2):
            table.cell(1, 3).text = re.sub(degreechange, "（B）", table.cell(1, 3).text)
        elif (auth.degree == 3):
            table.cell(1, 3).text = re.sub(degreechange, "（C）", table.cell(1, 3).text)
        elif (auth.degree == 4):
            table.cell(1, 3).text = re.sub(degreechange, "（D）", table.cell(1, 3).text)
        # 第一作者
        table.cell(2, 3).text = auth.major
        table.cell(2, 8).text = auth.enroll_time.strftime("%Y-%m-%d")
        table.cell(3, 4).text = work_info.title
        table.cell(4, 3).text = auth.address
        table.cell(4, 8).text = auth.phone
        table.cell(5, 8).text = auth.user.email
        # 其他合作者
        row = 7
        for i in company:
            table.cell(row, 1).text = i.name
            table.cell(row, 2).text = i.stu_id
            if (i.degree == 1):
                table.cell(row, 3).text = "大专"
            elif (i.degree == 2):
                table.cell(row, 3).text = "大学本科"
            elif (i.degree == 3):
                table.cell(row, 3).text = "硕士研究生"
            elif (i.degree == 4):
                table.cell(row, 3).text = "博士研究生"
            table.cell(row, 6).text = i.phone
            table.cell(row, 7).text = i.user.email
            row = row + 1
        # 表格2
        table2 = tables[1]  # 获取文件中的第二个表格
        table2.cell(0, 1).text = work_info.title
        fieldchange = r"（  ）"
        if (work_info.field == 1):
            table2.cell(1, 1).text = re.sub(fieldchange, "（A）", table2.cell(1, 1).text)
        elif (work_info.field == 2):
            table2.cell(1, 1).text = re.sub(fieldchange, "（B）", table2.cell(1, 1).text)
        elif (work_info.field == 3):
            table2.cell(1, 1).text = re.sub(fieldchange, "（C）", table2.cell(1, 1).text)
        elif (work_info.field == 4):
            table2.cell(1, 1).text = re.sub(fieldchange, "（D）", table2.cell(1, 1).text)
        elif (work_info.field == 5):
            table2.cell(1, 1).text = re.sub(fieldchange, "（E）", table2.cell(1, 1).text)
        elif (work_info.field == 6):
            table2.cell(1, 1).text = re.sub(fieldchange, "（F）", table2.cell(1, 1).text)

        table2.cell(2, 1).text = work_info.detail
        table2.cell(3, 1).text = work_info.innovation
        table2.cell(4, 1).text = work_info.keywords
        labellist = json.loads(work_info.labels)['labels']
        labellist=list(map(int,labellist))
        labellist.sort(reverse=True)

        position = ()
        if (work_info.work_type == 1):
            position = (5, 1)
        elif (work_info.work_type == 2):
            position = (6, 1)
        opt_list = table2.cell(*position).text.split("□")[1:]
        num = len(opt_list)
        checklist = ["□" for i in range(num)]
        for i in labellist:
            checklist[i - 1] = "☑"
        check_result = ""
        for check, opt in zip(checklist, opt_list):
            check_result += check + opt
        table2.cell(*position).text = check_result
        document.save('static/pdf/fulltable.docx')
        convertapi.api_secret = 'D1kgtEI0Qc5VTJpb'
        convertapi.convert('pdf', {
            'File': 'static/pdf/fulltable.docx'
        }, from_format='docx').save_files('static/pdf/fulltable.pdf')


    return JsonResponse({'ret': 'Reply for work_id'})

